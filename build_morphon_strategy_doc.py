from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "MORPHON_Executive_Summary_and_Sellable_Offers.docx"

COLORS = {
    "graphite": RGBColor(11, 15, 18),
    "charcoal": RGBColor(31, 38, 44),
    "muted": RGBColor(101, 112, 122),
    "line": RGBColor(210, 218, 222),
    "accent": RGBColor(0, 166, 184),
    "soft": RGBColor(242, 246, 247),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D7DEE2", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="00A6B8", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLORS["charcoal"]
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 26, COLORS["graphite"], 0, 8),
        ("Subtitle", 11, COLORS["muted"], 0, 18),
        ("Heading 1", 16, COLORS["graphite"], 16, 7),
        ("Heading 2", 13, COLORS["graphite"], 11, 5),
        ("Heading 3", 11, COLORS["graphite"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "MORPHON | Executive Summary & Sellable Offers"
    p.style = "Normal"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = COLORS["muted"]
    add_bottom_border(p, color="D7DEE2", size="4", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["muted"]


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = COLORS["accent"]
    p.paragraph_format.space_after = Pt(5)


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        r.font.color.rgb = COLORS["graphite"]
        p.add_run(text[len(bold_start) :])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F6F7")
    set_cell_borders(cell, "CAD8DD", "5")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["graphite"]
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.add_run(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "EAF3F5")
        set_cell_borders(hdr[i])
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLORS["graphite"]
                r.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_borders(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9.2)
                    r.font.color.rgb = COLORS["charcoal"]
    doc.add_paragraph()
    return table


def add_offer(doc, number, name, description, gets, best_for, promise):
    doc.add_heading(f"{number}. {name}", level=2)
    add_body(doc, description)
    add_table(
        doc,
        ["What the client gets", "Best for", "Sellable promise"],
        [["\n".join(gets), "\n".join(best_for), promise]],
        widths=[2.35, 2.05, 2.1],
    )


def build():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_header_footer(section)

    add_kicker(doc, "Company strategy document")
    title = doc.add_paragraph(style="Title")
    title.add_run("MORPHON")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Executive Summary & Sellable Offers\nParametric design-to-build systems for architecture, engineering, construction, and fabrication teams.")
    add_bottom_border(subtitle, color="00A6B8", size="10", space="10")
    add_body(
        doc,
        "MORPHON develops computational design systems that transform complex architectural, structural, and environmental problems into intelligent parametric models, BIM workflows, analysis tools, documentation packages, and fabrication-ready outputs.",
    )
    add_callout(
        doc,
        "Core promise",
        "We help AEC teams transform complex projects into intelligent models that can be designed, analyzed, documented, quantified, and fabricated with greater control.",
    )

    doc.add_heading("Executive Summary", level=1)
    add_body(
        doc,
        "MORPHON is a computational design and AEC technology company positioned between architecture, engineering, construction, computation, and fabrication. It is not a traditional architecture studio, a generic BIM outsourcing service, or only a software company.",
    )
    add_body(
        doc,
        "The company creates living project systems: models and workflows that can change, update, calculate, document, and support decisions across the project lifecycle. Instead of producing isolated drawings, static 3D models, or disconnected analyses, MORPHON connects geometry, data, performance, documentation, and fabrication logic.",
    )
    add_body(
        doc,
        "The immediate business should be simple and focused: MORPHON sells high-value computational design systems that help clients design faster, reduce uncertainty, and build complex projects with more control.",
    )
    doc.add_heading("Primary Capabilities", level=2)
    add_bullets(
        doc,
        [
            "Complex geometry and geometric rationalization",
            "Parametric BIM and documentation automation",
            "Tensile, lightweight, facade, canopy, and envelope systems",
            "Environmental and structural-aware design analysis",
            "Fabrication logic, shop drawings, part systems, and CNC-ready outputs",
            "Custom digital workflows, web configurators, and presales tools",
        ],
    )

    doc.add_heading("Positioning", level=1)
    add_table(
        doc,
        ["Version", "Positioning statement"],
        [
            ["One-line", "MORPHON creates parametric design and BIM systems for complex architecture, engineering, and fabrication."],
            ["Premium", "MORPHON helps AEC teams turn complex design intent into intelligent parametric models, performance analysis, construction documentation, and fabrication-ready systems."],
            ["Landing page", "From complex geometry to buildable systems. Parametric design, BIM automation, environmental analysis, and fabrication-ready documentation for ambitious AEC projects."],
        ],
        widths=[1.35, 5.15],
    )

    doc.add_heading("The Core Problem", level=1)
    add_body(doc, "Most AEC projects suffer from fragmentation.")
    add_bullets(
        doc,
        [
            "Design happens in one place, engineering in another, BIM somewhere else.",
            "Analysis is separate, fabrication comes late, and documentation is slow.",
            "Changes create manual rework, coordination risk, and uncertainty.",
        ],
    )
    add_body(
        doc,
        "MORPHON solves this by creating parametric project systems where geometry, performance, data, and documentation are connected. When the design changes, the system can update.",
    )
    add_bullets(
        doc,
        [
            "Faster iteration and clearer design options",
            "Better technical control and fewer disconnected models",
            "Less manual rework across documentation and coordination",
            "More buildable complexity and earlier issue detection",
            "Clearer communication with clients, engineers, fabricators, and contractors",
        ],
    )

    doc.add_heading("Offer Architecture", level=1)
    add_body(
        doc,
        "MORPHON should not sell Grasshopper, Rhino, BIM, simulations, or parametric design as isolated technical services. The company should sell business outcomes: speed, control, buildability, coordination, and market-ready digital tools.",
    )
    add_table(
        doc,
        ["Service line", "What it includes", "Landing-page headline"],
        [
            [
                "Parametric Design Systems",
                "Complex geometry, parametric facades, tensile structures, structural-aware design, environmental analysis.",
                "Design systems for complex architecture.",
            ],
            [
                "BIM + Documentation Automation",
                "Parametric BIM, Revit/Rhino workflows, construction documentation, quantity takeoffs, shop drawings, fabrication models.",
                "From model to documentation with less manual rework.",
            ],
            [
                "Digital Products + Configurators",
                "Online configurators, 3D product visualization, sales tools, automated quoting, internal AEC tools, digital twins.",
                "Interactive tools for selling and managing AEC products.",
            ],
        ],
        widths=[1.75, 3.05, 1.7],
    )

    doc.add_heading("Flagship Offer", level=1)
    add_callout(
        doc,
        "Parametric Design-to-Build System",
        "MORPHON creates custom parametric systems that take complex architectural or structural ideas from concept to buildable reality: design exploration, BIM modeling, performance analysis, documentation, quantification, and fabrication logic.",
    )
    doc.add_heading("Why this offer is strong", level=2)
    add_bullets(
        doc,
        [
            "It combines the company's capabilities without sounding fragmented.",
            "It avoids selling tools as the product.",
            "It communicates a business outcome: delivering complex projects with speed, intelligence, and control.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Sellable Offers", level=1)

    offers = [
        (
            "Parametric BIM Delivery System",
            "We create intelligent BIM models driven by parametric logic, allowing projects to be designed, modified, quantified, documented, and coordinated with greater speed and control.",
            ["Parametric 3D/BIM model", "Flexible design parameters", "Automated geometry updates", "Quantity extraction", "Construction documentation", "Optional Revit/Rhino/Grasshopper integration"],
            ["Architecture firms", "Developers", "Engineering offices", "Construction companies", "Projects with many variable elements"],
            "Turn your project into an intelligent model that updates with design changes and supports documentation, coordination, and construction decisions.",
        ),
        (
            "Complex Geometry Design System",
            "We help teams design, rationalize, and document complex architectural geometries that are difficult to control with conventional CAD/BIM workflows.",
            ["Parametric geometry system", "Design variation studies", "Geometry rationalization", "Surface, panel, strip, or module logic", "Buildability strategy", "Fabrication-ready geometry if needed"],
            ["Freeform roofs", "Sculptural facades", "Canopies", "Pavilions", "Shells", "Installations", "Organic or non-standard architecture"],
            "Turn complex forms into controllable, rationalized, and buildable systems.",
        ),
        (
            "Parametric Facade System",
            "We design facade systems where geometry, modules, openings, structure, environmental performance, and documentation are connected parametrically.",
            ["Parametric facade model", "Panelization logic", "Module variation system", "Solar and shadow analysis", "Documentation drawings", "Fabrication or installation logic", "Quantity takeoffs"],
            ["Developers", "Architecture firms", "Facade consultants", "Commercial buildings", "Mixed-use and institutional projects"],
            "Design smarter facades that connect geometry, climate, cost, and construction logic from the beginning.",
        ),
        (
            "Tensile Structure Design + Fabrication System",
            "We create parametric workflows for tensile structures, canopies, membranes, and lightweight systems from initial form-finding to fabrication support.",
            ["Initial form-finding model", "Parametric membrane or cable geometry", "Structural-aware iterations", "Patterning logic", "Connection studies", "Shop drawings", "Fabrication geometry"],
            ["Tensile structure companies", "Fabricators", "Architecture firms", "Lightweight roofs", "Sports, cultural, commercial, and public-space projects"],
            "Move from conceptual tensile form to fabrication-ready system with parametric control.",
        ),
        (
            "Structural-Aware Design System",
            "We help design teams generate and test forms that are visually ambitious and informed by structural behavior, fabrication logic, and geometric performance.",
            ["Parametric structural model", "Iterative design options", "Preliminary structural analysis", "Karamba/Grasshopper workflows", "Optimization strategy", "Engineering coordination package"],
            ["Shells", "Thin structures", "Canopies", "Pavilions", "Complex roofs", "Experimental architecture"],
            "Explore ambitious forms with structural intelligence from the earliest design stages.",
        ),
        (
            "Climate-Aware Design Analysis",
            "We integrate environmental data into the design process so teams can understand how sunlight, radiation, shadows, airflow, water, and site conditions affect the project.",
            ["Solar radiation analysis", "Shadow studies", "Daylight studies", "Wind or airflow analysis when relevant", "Water flow or drainage logic", "Comparative scenarios", "Visual reports"],
            ["Architecture firms", "Urban designers", "Developers", "Public-space projects", "Climate-sensitive buildings"],
            "Make design decisions using environmental intelligence, not intuition alone.",
        ),
        (
            "Design-to-Fabrication System",
            "We convert complex designs into organized, documented, and fabrication-ready systems for manufacturing, assembly, and construction.",
            ["Fabrication model", "Part numbering", "Assembly logic", "CNC-ready files depending on material", "Shop drawings", "Installation drawings", "Quantity schedules"],
            ["Fabricators", "Metal workshops", "Facade contractors", "Tensile structure companies", "Custom architecture elements"],
            "Translate complex design into parts, drawings, files, and assembly logic that can actually be built.",
        ),
        (
            "AEC Product Configurator",
            "We create interactive online configurators that allow clients to customize, price, visualize, and presell architectural or construction products.",
            ["Web-based 3D configurator", "Parametric product logic", "Real-time visualization", "Option selection", "Pricing or quantity logic", "Lead-generation system", "Optional automated quote generation"],
            ["Steel structure companies", "Pergola and canopy manufacturers", "Facade product companies", "Modular construction companies", "Architectural product brands"],
            "Turn your construction product into an interactive sales tool that helps clients visualize, customize, and request quotes online.",
        ),
        (
            "AEC Workflow Automation",
            "We automate repetitive design, modeling, documentation, and coordination tasks using custom computational workflows.",
            ["Grasshopper, Python, C#, Revit, and Rhino workflows", "Automation scripts", "Parametric tools", "BIM data workflows", "Drawing automation", "Quantity automation", "Team documentation"],
            ["Architecture offices", "BIM teams", "Engineering offices", "Fabricators", "Construction companies"],
            "Reduce weeks of repetitive technical work into hours or days through custom computational workflows.",
        ),
    ]
    for i, offer in enumerate(offers, 1):
        add_offer(doc, i, *offer)

    doc.add_heading("Landing Page Direction", level=1)
    add_table(
        doc,
        ["Element", "Recommended direction"],
        [
            ["Hero headline", "Complex architecture, made buildable."],
            ["Hero copy", "MORPHON creates parametric design, BIM, analysis, and fabrication systems for architecture, engineering, and construction teams."],
            ["CTA language", "Start a Project / Explore Services"],
            ["Visual language", "Parametric grids, structural meshes, facade panels, tensile membranes, nodal systems, and precise technical diagrams."],
            ["Tone", "Premium, precise, technical, clear, and outcome-driven."],
            ["Palette", "Graphite, charcoal, off-white, muted steel, controlled cyan, and a sparing acid-green or warm-metal accent."],
            ["Typography", "Space Grotesk or Geist for headings; Inter or Geist Sans for body; IBM Plex Mono or Geist Mono for technical labels."],
        ],
        widths=[1.55, 4.95],
    )
    add_callout(
        doc,
        "Cleanest public-facing offer",
        "MORPHON develops parametric design-to-build systems for AEC teams working with complex geometry, BIM automation, environmental analysis, and fabrication.",
    )

    doc.core_properties.title = "MORPHON Executive Summary and Sellable Offers"
    doc.core_properties.subject = "Company positioning, offer architecture, and landing page direction"
    doc.core_properties.author = "MORPHON"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
