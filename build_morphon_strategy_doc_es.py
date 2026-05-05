import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "MORPHON_Resumen_Ejecutivo_y_Ofertas_Comerciales.docx"

COLORS = {
    "graphite": RGBColor(11, 15, 18),
    "charcoal": RGBColor(31, 38, 44),
    "muted": RGBColor(101, 112, 122),
    "accent": RGBColor(0, 166, 184),
}

SPANISH_FIXES = {
    "entender como": "entender cómo",
    "parametricos": "paramétricos",
    "parametrico": "paramétrico",
    "parametricas": "paramétricas",
    "parametrica": "paramétrica",
    "parametricamente": "paramétricamente",
    "geometrica": "geométrica",
    "geometricas": "geométricas",
    "geometrico": "geométrico",
    "geometricos": "geométricos",
    "geometria": "geometría",
    "geometrias": "geometrías",
    "diseno": "diseño",
    "disenos": "diseños",
    "disenar": "diseñar",
    "disenarse": "diseñarse",
    "disenamos": "diseñamos",
    "Disenamos": "Diseñamos",
    "disena": "diseña",
    "Disena": "Diseña",
    "arquitectonicos": "arquitectónicos",
    "arquitectonicas": "arquitectónicas",
    "tecnico": "técnico",
    "tecnica": "técnica",
    "tecnicos": "técnicos",
    "tecnicas": "técnicas",
    "tecnologia": "tecnología",
    "computacion": "computación",
    "ingenieria": "ingeniería",
    "analisis": "análisis",
    "mas": "más",
    "estan": "están",
    "esta separado": "está separado",
    "solo": "solo",
    "logica": "lógica",
    "logicas": "lógicas",
    "documentacion": "documentación",
    "automatizacion": "automatización",
    "fabricacion": "fabricación",
    "construccion": "construcción",
    "coordinacion": "coordinación",
    "cuantificacion": "cuantificación",
    "cuantificaciones": "cuantificaciones",
    "iteracion": "iteración",
    "actualizaciones automaticas": "actualizaciones automáticas",
    "automaticas": "automáticas",
    "extraccion": "extracción",
    "integracion": "integración",
    "cotizacion": "cotización",
    "generacion": "generación",
    "seleccion": "selección",
    "visualizacion": "visualización",
    "direccion": "dirección",
    "Direccion": "Dirección",
    "Version": "Versión",
    "Declaracion": "Declaración",
    "Linea": "Línea",
    "Que incluye": "Qué incluye",
    "Por que": "Por qué",
    "dificiles": "difíciles",
    "mayoria": "mayoría",
    "fragmentacion": "fragmentación",
    "rapida": "rápida",
    "deteccion": "detección",
    "comunicacion": "comunicación",
    "racionalizacion": "racionalización",
    "parametros": "parámetros",
    "variacion": "variación",
    "escultoricas": "escultóricas",
    "panelizacion": "panelización",
    "instalacion": "instalación",
    "optimizacion": "optimización",
    "radiacion": "radiación",
    "intuicion": "intuición",
    "seleccion": "selección",
    "reticulas": "retículas",
    "calido": "cálido",
    "minimo": "mínimo",
    "acido": "ácido",
    "metalico": "metálico",
    "tipografia": "tipografía",
    "exploracion": "exploración",
    "publica": "pública",
    "a medida": "a medida",
    "desempeno": "desempeño",
    "rapido": "rápido",
    "generico": "genérico",
    "estaticos": "estáticos",
    "ciclo completo": "ciclo completo",
    "modulos": "módulos",
    "modular": "modular",
    "organica": "orgánica",
    "estandar": "estándar",
    "metalicos": "metálicos",
    "pergolas": "pérgolas",
    "publico": "público",
    "climatico": "climático",
    "clima": "clima",
    "carbon": "carbón",
    "titulos": "títulos",
    "fisico": "físico",
    "linea": "línea",
    "energia": "energía",
    "numero": "número",
    "Numeracion": "Numeración",
    "segun": "según",
    "dias": "días",
}


def fix_spanish_text(text):
    fixed = text
    # Apply longer phrases first so phrase-level fixes win before individual words.
    for raw, replacement in sorted(SPANISH_FIXES.items(), key=lambda item: len(item[0]), reverse=True):
        def preserve_case(match):
            value = match.group(0)
            if value.isupper():
                return replacement.upper()
            if value[:1].isupper():
                return replacement[:1].upper() + replacement[1:]
            return replacement

        fixed = re.sub(rf"(?<!\w){re.escape(raw)}(?!\w)", preserve_case, fixed, flags=re.IGNORECASE)
    return fixed


def normalize_doc_text(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = fix_spanish_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = fix_spanish_text(run.text)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    run.text = fix_spanish_text(run.text)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D7DEE2", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="00A6B8", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = COLORS["charcoal"]
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 26, COLORS["graphite"], 0, 8),
        ("Subtitle", 11, COLORS["muted"], 0, 18),
        ("Heading 1", 16, COLORS["graphite"], 16, 7),
        ("Heading 2", 13, COLORS["graphite"], 11, 5),
        ("Heading 3", 11, COLORS["graphite"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(section):
    p = section.header.paragraphs[0]
    p.text = "MORPHON | Resumen Ejecutivo y Ofertas Comerciales"
    p.style = "Normal"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = COLORS["muted"]
    add_bottom_border(p, color="D7DEE2", size="4", space="4")

    p = section.footer.paragraphs[0]
    add_page_number(p)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["muted"]


def add_kicker(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = COLORS["accent"]
    p.paragraph_format.space_after = Pt(5)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F6F7")
    set_cell_borders(cell, "CAD8DD", "5")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["graphite"]
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.add_run(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "EAF3F5")
        set_cell_borders(hdr[i])
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLORS["graphite"]
                r.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_borders(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9.2)
                    r.font.color.rgb = COLORS["charcoal"]
    doc.add_paragraph()
    return table


def add_offer(doc, number, name, description, gets, best_for, promise):
    doc.add_heading(f"{number}. {name}", level=2)
    add_body(doc, description)
    add_table(
        doc,
        ["Lo que recibe el cliente", "Ideal para", "Promesa comercial"],
        [["\n".join(gets), "\n".join(best_for), promise]],
        widths=[2.35, 2.05, 2.1],
    )


def build():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_header_footer(section)

    add_kicker(doc, "Documento de estrategia de empresa")
    title = doc.add_paragraph(style="Title")
    title.add_run("MORPHON")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Resumen Ejecutivo y Ofertas Comerciales\nSistemas parametricos de diseno-a-construccion para arquitectura, ingenieria, construccion y fabricacion.")
    add_bottom_border(subtitle, color="00A6B8", size="10", space="10")
    add_body(
        doc,
        "MORPHON desarrolla sistemas de diseno computacional que transforman problemas arquitectonicos, estructurales y ambientales complejos en modelos parametricos inteligentes, flujos BIM, herramientas de analisis, paquetes de documentacion y salidas listas para fabricacion.",
    )
    add_callout(
        doc,
        "Promesa central",
        "Ayudamos a equipos AEC a transformar proyectos complejos en modelos inteligentes que pueden disenarse, analizarse, documentarse, cuantificarse y fabricarse con mayor control.",
    )

    doc.add_heading("Resumen Ejecutivo", level=1)
    add_body(
        doc,
        "MORPHON es una empresa de diseno computacional y tecnologia AEC ubicada entre arquitectura, ingenieria, construccion, computacion y fabricacion. No es un despacho tradicional de arquitectura, un servicio generico de outsourcing BIM ni solamente una empresa de software.",
    )
    add_body(
        doc,
        "La empresa crea sistemas vivos de proyecto: modelos y flujos de trabajo que pueden cambiar, actualizarse, calcular, documentar y apoyar decisiones durante el ciclo completo del proyecto. En lugar de producir dibujos aislados, modelos 3D estaticos o analisis desconectados, MORPHON conecta geometria, datos, desempeno, documentacion y logica de fabricacion.",
    )
    add_body(
        doc,
        "El negocio inmediato debe ser simple y enfocado: MORPHON vende sistemas computacionales de alto valor que ayudan a los clientes a disenar mas rapido, reducir incertidumbre y construir proyectos complejos con mas control.",
    )
    doc.add_heading("Capacidades Principales", level=2)
    add_bullets(
        doc,
        [
            "Geometria compleja y racionalizacion geometrica",
            "BIM parametrico y automatizacion de documentacion",
            "Sistemas tensiles, ligeros, fachadas, cubiertas y envolventes",
            "Analisis ambiental y diseno informado por comportamiento estructural",
            "Logica de fabricacion, planos de taller, sistemas de piezas y salidas CNC",
            "Flujos digitales a medida, configuradores web y herramientas de preventa",
        ],
    )

    doc.add_heading("Posicionamiento", level=1)
    add_table(
        doc,
        ["Version", "Declaracion de posicionamiento"],
        [
            ["Una linea", "MORPHON crea sistemas de diseno parametrico y BIM para arquitectura, ingenieria y fabricacion compleja."],
            ["Premium", "MORPHON ayuda a equipos AEC a convertir intenciones de diseno complejas en modelos parametricos inteligentes, analisis de desempeno, documentacion constructiva y sistemas listos para fabricacion."],
            ["Landing page", "De geometria compleja a sistemas construibles. Diseno parametrico, automatizacion BIM, analisis ambiental y documentacion lista para fabricacion para proyectos AEC ambiciosos."],
        ],
        widths=[1.35, 5.15],
    )

    doc.add_heading("El Problema Central", level=1)
    add_body(doc, "La mayoria de los proyectos AEC sufren por fragmentacion.")
    add_bullets(
        doc,
        [
            "El diseno ocurre en un lugar, la ingenieria en otro y el BIM en otro mas.",
            "El analisis esta separado, la fabricacion llega tarde y la documentacion avanza lentamente.",
            "Los cambios generan retrabajo manual, riesgos de coordinacion e incertidumbre.",
        ],
    )
    add_body(
        doc,
        "MORPHON resuelve esto mediante sistemas parametricos de proyecto donde geometria, desempeno, datos y documentacion estan conectados. Cuando el diseno cambia, el sistema puede actualizarse.",
    )
    add_bullets(
        doc,
        [
            "Iteracion mas rapida y opciones de diseno mas claras",
            "Mayor control tecnico y menos modelos desconectados",
            "Menos retrabajo manual en documentacion y coordinacion",
            "Complejidad mas construible y deteccion temprana de problemas",
            "Comunicacion mas clara con clientes, ingenieros, fabricadores y contratistas",
        ],
    )

    doc.add_heading("Arquitectura de Ofertas", level=1)
    add_body(
        doc,
        "MORPHON no debe vender Grasshopper, Rhino, BIM, simulaciones o diseno parametrico como servicios tecnicos aislados. La empresa debe vender resultados de negocio: velocidad, control, construibilidad, coordinacion y herramientas digitales listas para el mercado.",
    )
    add_table(
        doc,
        ["Linea de servicio", "Que incluye", "Headline para landing page"],
        [
            [
                "Sistemas de Diseno Parametrico",
                "Geometria compleja, fachadas parametricas, estructuras tensiles, diseno estructuralmente informado y analisis ambiental.",
                "Sistemas de diseno para arquitectura compleja.",
            ],
            [
                "BIM + Automatizacion de Documentacion",
                "BIM parametrico, flujos Revit/Rhino, documentacion constructiva, cuantificaciones, planos de taller y modelos de fabricacion.",
                "Del modelo a la documentacion con menos retrabajo manual.",
            ],
            [
                "Productos Digitales + Configuradores",
                "Configuradores online, visualizacion 3D de productos, herramientas comerciales, cotizacion automatizada, herramientas internas AEC y gemelos digitales.",
                "Herramientas interactivas para vender y gestionar productos AEC.",
            ],
        ],
        widths=[1.75, 3.05, 1.7],
    )

    doc.add_heading("Oferta Insignia", level=1)
    add_callout(
        doc,
        "Sistema Parametrico de Diseno-a-Construccion",
        "MORPHON crea sistemas parametricos a medida que llevan ideas arquitectonicas o estructurales complejas desde el concepto hasta una realidad construible: exploracion de diseno, modelado BIM, analisis de desempeno, documentacion, cuantificacion y logica de fabricacion.",
    )
    doc.add_heading("Por que esta oferta es fuerte", level=2)
    add_bullets(
        doc,
        [
            "Integra las capacidades de la empresa sin sonar fragmentada.",
            "Evita vender herramientas como si fueran el producto.",
            "Comunica un resultado de negocio: entregar proyectos complejos con velocidad, inteligencia y control.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Ofertas Comerciales", level=1)

    offers = [
        (
            "Sistema de Entrega BIM Parametrico",
            "Creamos modelos BIM inteligentes impulsados por logica parametrica, permitiendo disenar, modificar, cuantificar, documentar y coordinar proyectos con mayor velocidad y control.",
            ["Modelo 3D/BIM parametrico", "Parametros flexibles de diseno", "Actualizaciones automaticas de geometria", "Extraccion de cantidades", "Documentacion constructiva", "Integracion opcional Revit/Rhino/Grasshopper"],
            ["Despachos de arquitectura", "Desarrolladores", "Oficinas de ingenieria", "Constructoras", "Proyectos con elementos variables"],
            "Convierte tu proyecto en un modelo inteligente que se actualiza con los cambios de diseno y apoya decisiones de documentacion, coordinacion y construccion.",
        ),
        (
            "Sistema de Diseno para Geometria Compleja",
            "Ayudamos a equipos a disenar, racionalizar y documentar geometrias arquitectonicas complejas dificiles de controlar con flujos CAD/BIM convencionales.",
            ["Sistema de geometria parametrica", "Estudios de variacion", "Racionalizacion geometrica", "Logica de superficies, paneles, tiras o modulos", "Estrategia de construibilidad", "Geometria lista para fabricacion si aplica"],
            ["Cubiertas libres", "Fachadas escultoricas", "Canopies", "Pabellones", "Cascarones", "Instalaciones", "Arquitectura organica o no estandar"],
            "Convierte formas complejas en sistemas controlables, racionalizados y construibles.",
        ),
        (
            "Sistema de Fachada Parametrica",
            "Disenamos sistemas de fachada donde geometria, modulos, aperturas, estructura, desempeno ambiental y documentacion estan conectados parametricamente.",
            ["Modelo parametrico de fachada", "Logica de panelizacion", "Sistema de variacion modular", "Analisis solar y de sombras", "Planos de documentacion", "Logica de fabricacion o instalacion", "Cuantificaciones"],
            ["Desarrolladores", "Despachos de arquitectura", "Consultores de fachada", "Edificios comerciales", "Proyectos mixtos e institucionales"],
            "Disena fachadas mas inteligentes conectando geometria, clima, costo y logica constructiva desde el inicio.",
        ),
        (
            "Sistema de Diseno + Fabricacion para Estructuras Tensiles",
            "Creamos flujos parametricos para estructuras tensiles, canopies, membranas y sistemas ligeros desde el form-finding inicial hasta el apoyo a fabricacion.",
            ["Modelo inicial de form-finding", "Geometria parametrica de membrana o cables", "Iteraciones estructuralmente informadas", "Logica de patronaje", "Estudios de conexiones", "Planos de taller", "Geometria de fabricacion"],
            ["Empresas de estructuras tensiles", "Fabricadores", "Despachos de arquitectura", "Cubiertas ligeras", "Proyectos deportivos, culturales, comerciales y de espacio publico"],
            "Pasa de una forma tensile conceptual a un sistema listo para fabricacion con control parametrico.",
        ),
        (
            "Sistema de Diseno Estructuralmente Informado",
            "Ayudamos a equipos de diseno a generar y probar formas visualmente ambiciosas informadas por comportamiento estructural, logica de fabricacion y desempeno geometrico.",
            ["Modelo estructural parametrico", "Opciones iterativas de diseno", "Analisis estructural preliminar", "Flujos Karamba/Grasshopper", "Estrategia de optimizacion", "Paquete de coordinacion con ingenieria"],
            ["Cascarones", "Estructuras delgadas", "Canopies", "Pabellones", "Cubiertas complejas", "Arquitectura experimental"],
            "Explora formas ambiciosas con inteligencia estructural desde las primeras etapas de diseno.",
        ),
        (
            "Analisis de Diseno Climatico",
            "Integramos datos ambientales al proceso de diseno para entender como sol, radiacion, sombras, flujo de aire, agua y condiciones del sitio afectan el proyecto.",
            ["Analisis de radiacion solar", "Estudios de sombra", "Estudios de luz natural", "Analisis de viento o flujo de aire cuando aplique", "Logica de escurrimiento o drenaje", "Escenarios comparativos", "Reportes visuales"],
            ["Despachos de arquitectura", "Urbanistas", "Desarrolladores", "Proyectos de espacio publico", "Edificios sensibles al clima"],
            "Toma decisiones de diseno con inteligencia ambiental, no solo con intuicion.",
        ),
        (
            "Sistema de Diseno-a-Fabricacion",
            "Convertimos disenos complejos en sistemas organizados, documentados y listos para fabricacion, manufactura, ensamble y construccion.",
            ["Modelo de fabricacion", "Numeracion de piezas", "Logica de ensamble", "Archivos CNC segun material", "Planos de taller", "Planos de instalacion", "Cuantificaciones"],
            ["Fabricadores", "Talleres metalicos", "Contratistas de fachada", "Empresas de estructuras tensiles", "Elementos arquitectonicos especiales"],
            "Traduce disenos complejos en piezas, planos, archivos y logica de ensamble que realmente pueden construirse.",
        ),
        (
            "Configurador de Producto AEC",
            "Creamos configuradores online interactivos que permiten a clientes personalizar, cotizar, visualizar y prevender productos arquitectonicos o de construccion.",
            ["Configurador 3D web", "Logica parametrica de producto", "Visualizacion en tiempo real", "Seleccion de opciones", "Logica de precios o cantidades", "Sistema de generacion de leads", "Cotizacion automatica opcional"],
            ["Empresas de estructura metalica", "Fabricantes de pergolas y canopies", "Empresas de productos de fachada", "Construccion modular", "Marcas de productos arquitectonicos"],
            "Convierte tu producto de construccion en una herramienta comercial interactiva para visualizar, personalizar y solicitar cotizaciones online.",
        ),
        (
            "Automatizacion de Flujos AEC",
            "Automatizamos tareas repetitivas de diseno, modelado, documentacion y coordinacion mediante flujos computacionales a medida.",
            ["Flujos Grasshopper, Python, C#, Revit y Rhino", "Scripts de automatizacion", "Herramientas parametricas", "Flujos de datos BIM", "Automatizacion de dibujos", "Automatizacion de cantidades", "Documentacion para equipos"],
            ["Despachos de arquitectura", "Equipos BIM", "Oficinas de ingenieria", "Fabricadores", "Constructoras"],
            "Reduce semanas de trabajo tecnico repetitivo a horas o dias mediante flujos computacionales personalizados.",
        ),
    ]
    for i, offer in enumerate(offers, 1):
        add_offer(doc, i, *offer)

    doc.add_heading("Direccion para Landing Page", level=1)
    add_table(
        doc,
        ["Elemento", "Direccion recomendada"],
        [
            ["Headline hero", "Arquitectura compleja, hecha construible."],
            ["Copy hero", "MORPHON crea sistemas de diseno parametrico, BIM, analisis y fabricacion para equipos de arquitectura, ingenieria y construccion."],
            ["CTA", "Iniciar un Proyecto / Explorar Servicios"],
            ["Lenguaje visual", "Reticulas parametricas, mallas estructurales, paneles de fachada, membranas tensiles, sistemas nodales y diagramas tecnicos precisos."],
            ["Tono", "Premium, preciso, tecnico, claro y orientado a resultados."],
            ["Paleta", "Grafito, carbon, blanco calido, acero tenue, cyan controlado y un acento minimo verde acido o metalico calido."],
            ["Tipografia", "Space Grotesk o Geist para titulos; Inter o Geist Sans para cuerpo; IBM Plex Mono o Geist Mono para etiquetas tecnicas."],
        ],
        widths=[1.55, 4.95],
    )
    add_callout(
        doc,
        "Oferta publica mas limpia",
        "MORPHON desarrolla sistemas parametricos de diseno-a-construccion para equipos AEC que trabajan con geometria compleja, automatizacion BIM, analisis ambiental y fabricacion.",
    )

    doc.core_properties.title = "MORPHON Resumen Ejecutivo y Ofertas Comerciales"
    doc.core_properties.subject = "Posicionamiento de empresa, arquitectura de ofertas y direccion para landing page"
    doc.core_properties.author = "MORPHON"
    normalize_doc_text(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
